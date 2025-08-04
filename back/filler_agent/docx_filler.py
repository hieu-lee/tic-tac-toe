"""
DOCX form filling functionality.

This module handles filling of DOCX forms, including text replacement,
checkbox updates, and font management.
"""

import logging
from pathlib import Path
from typing import List, Optional, Literal
from docx import Document
from .fill_processor import detect_fill_entries, process_fill_entries, FillEntry
from .checkbox_processor import (
    CheckboxEntry,
    detect_checkbox_entries,
    process_checkbox_entries,
    update_checkbox_in_paragraph,
)
from .text_utils import replace_text_preserve_layout_docx


def fill_docx_with_entries(
    fill_entries: List[FillEntry],
    checkbox_entries: List[CheckboxEntry],
    form_path: str,
    output_path: Optional[str] = None,
) -> str:
    """Fill a DOCX form by performing a layout-preserving find-and-replace for every line in each entry.

    The replacement strategy uses the new replace_text_preserve_layout_docx function which:
    1. Preserves formatting and styling of the original text
    2. Handles multi-line replacements correctly
    3. Maintains proper ordering when multiple identical strings exist
    """

    if not output_path:
        output_path = form_path.replace(".docx", "_filled.docx")

    form_path_p = Path(form_path)
    temp_path = form_path_p.with_stem(form_path_p.stem + "_temp").with_suffix(".docx")

    s1_lines = []
    s2_lines = []

    for entry in fill_entries:
        s1_lines.extend(entry.lines.splitlines())
        s2_lines.extend(entry.filled_lines.splitlines())

    current_path = form_path

    if fill_entries:
        after_fill_path = str(temp_path) if checkbox_entries else output_path

        replace_text_preserve_layout_docx(
            current_path, s1_lines, s2_lines, out_path=after_fill_path
        )

        current_path = after_fill_path

    if checkbox_entries:
        s1_lines = []
        s2_lines = []
        for entry in checkbox_entries:
            s1_lines.extend(entry.lines.splitlines())
            s2_lines.extend(update_checkbox_in_paragraph(entry).splitlines())

        replace_text_preserve_layout_docx(
            current_path,
            s1_lines,
            s2_lines,
            out_path=output_path,
        )

    # Clean up temporary file if it exists
    if temp_path.exists():
        try:
            temp_path.unlink()
        except Exception as e:
            logging.warning(f"Failed to delete temporary file {temp_path}: {e}")

    return output_path


# Rewrite original fill_docx to preserve legacy behaviour but delegate to new implementation


def fill_docx(
    form_path: str,
    context_dir: str,
    output_path: Optional[str],
    placeholder_pattern,
    provider: Literal["openai", "groq", "google", "anythingllm", "local", "ollama"],
):
    """Legacy DOCX fill. Detects entries and checkboxes and delegates to fill_docx_with_entries."""
    doc = Document(form_path)

    # Extract all lines and track locations with original font information
    lines: List[str] = []
    locations = (
        []
    )  # tuples for replacing: ('para', paragraph, original_font_info) or ('cell', cell, paragraph, original_font_info)

    for para in doc.paragraphs:
        lines.append(para.text)
        # Extract font info from the first run with font information
        original_font_info = None
        for run in para.runs:
            if run.font.name:
                original_font_info = {
                    "name": run.font.name,
                    "size": run.font.size,
                    "bold": run.font.bold,
                    "italic": run.font.italic,
                    "underline": run.font.underline,
                }
                break
        locations.append(("para", para, original_font_info))

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    lines.append(para.text)
                    # Extract font info from the first run with font information
                    original_font_info = None
                    for run in para.runs:
                        if run.font.name:
                            original_font_info = {
                                "name": run.font.name,
                                "size": run.font.size,
                                "bold": run.font.bold,
                                "italic": run.font.italic,
                                "underline": run.font.underline,
                            }
                            break
                    locations.append(("cell", cell, para, original_font_info))

    # Detect fill entries
    entries = detect_fill_entries(lines, placeholder_pattern)
    # Process each entry: infer keys, fill missing context, generate filled_lines
    entries, _ = process_fill_entries(
        entries, context_dir, form_path, placeholder_pattern, provider
    )

    # # Detect and process checkbox entries
    checkbox_entries = detect_checkbox_entries(lines)
    checkbox_entries = process_checkbox_entries(
        checkbox_entries, context_dir, provider
    )

    return fill_docx_with_entries(entries, checkbox_entries, form_path, output_path)
