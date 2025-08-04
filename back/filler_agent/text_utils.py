"""
Text processing utilities for form filling.

This module contains utility functions for text processing, including
Unicode character sanitization for PDF rendering.
"""

from __future__ import annotations
import hashlib
import logging
import os

from pypdf import PdfReader

# Use font management utilities from font_manager
if __name__ == "__main__":
    # When running as a script, use direct import
    from font_manager import get_available_font
else:
    # When imported as a module, use relative import
    from .font_manager import get_available_font

from collections import Counter
import shutil
import tempfile
from pathlib import Path
from typing import Iterable

import fitz  # PyMuPDF
from docx import Document
from docx.text.paragraph import Paragraph


def is_interactive_pdf(form_path: str) -> bool:
    """Return **True** if *form_path* points to a PDF that contains at least one
    interactive AcroForm field (i.e. text widgets, checkboxes, etc.).

    The function is intentionally lightweight – it only parses the PDF catalog
    to inspect the *AcroForm* dictionary via *pypdf*'s ``PdfReader``. If any
    field objects are present the file is considered *interactive*.

    A ``False`` fallback is returned for any IO or parsing failure so that the
    caller can gracefully default to *flat* PDF handling.
    """

    try:
        reader = PdfReader(form_path)
        return bool(reader.get_fields())
    except Exception as exc:
        # Log at *debug* level to avoid noisy warnings for non-PDF inputs or
        # malformed documents while still providing insight during debugging.
        logging.debug(f"Interactive PDF detection failed for '{form_path}': {exc}")
        return False


def sanitize_unicode_for_pdf(text: str) -> str:
    """
    Sanitize Unicode characters that may not render properly in PDF fonts.
    Replaces problematic Unicode characters with ASCII equivalents.
    """
    if not text:
        return text

    # Track if any replacements were made
    original_text = text

    # Unicode character mappings to ASCII equivalents
    unicode_replacements = {
        # Smart quotes
        "\u201c": '"',  # Left double quotation mark (8220)
        "\u201d": '"',  # Right double quotation mark (8221)
        "\u2018": "'",  # Left single quotation mark (8216)
        "\u2019": "'",  # Right single quotation mark (8217)
        # Dashes
        "\u2014": "--",  # Em dash (8212)
        "\u2013": "-",  # En dash (8211)
        # Other common problematic characters
        "\u2026": "...",  # Horizontal ellipsis (8230)
        "\u00a0": " ",  # Non-breaking space (160)
    }

    # Apply replacements
    for unicode_char, ascii_replacement in unicode_replacements.items():
        text = text.replace(unicode_char, ascii_replacement)

    # Log if any replacements were made
    if text != original_text:
        replaced_chars = []
        for unicode_char, ascii_replacement in unicode_replacements.items():
            if unicode_char in original_text:
                replaced_chars.append(f"'{unicode_char}' → '{ascii_replacement}'")
        if replaced_chars:
            logging.debug(
                f"Sanitized Unicode characters for PDF rendering: {', '.join(replaced_chars)}"
            )

    return text


# ---------------------------------------------------------------------------
# Glyph model helpers
# ---------------------------------------------------------------------------


def _span_glyphs(span: dict) -> Iterable[dict]:
    """Yield glyph dictionaries for *span* compatible with PyMuPDF 1.22+."""
    # Raw dict format has a 'chars' list with per‑glyph info in MuPDF ≥ 1.22.
    if "chars" in span:
        for ch in span["chars"]:
            yield {
                "c": ch["c"],
                "bbox": ch["bbox"],
                "font": span["font"],
                "size": span["size"],
                "color": span["color"],
                "flags": span.get("flags", 0),
            }
    else:
        # Older PyMuPDF versions return only the whole text in span["text"].
        txt = span.get("text", "")
        w = fitz.Font(fontname=span["font"]).text_length
        x0, y0, x1, y1 = span["bbox"]
        step = (x1 - x0) / max(len(txt), 1)
        for idx, ch in enumerate(txt):
            yield {
                "c": ch,
                "bbox": (x0 + idx * step, y0, x0 + (idx + 1) * step, y1),
                "font": span["font"],
                "size": span["size"],
                "color": span["color"],
                "flags": span.get("flags", 0),
            }


def _normalize_color(color):
    """Return a color tuple that PyMuPDF accepts (components between 0 and 1).

    The *color* extracted from glyph dictionaries can take various forms,
    most commonly:

    1. An ``int`` packed as ``0xRRGGBB`` (legacy behaviour of ``get_text``).
    2. A tuple / list of 3 or 4 ``int``\s in the 0-255 range.
    3. A tuple / list of 3 or 4 ``float``\s already in the 0-1 range.

    This helper converts all variants to the canonical ``(r, g, b)`` tuple
    with float components between 0 and 1 required by drawing operators like
    ``page.draw_line`` and by ``TextWriter.color``.
    """

    # Already a 3-tuple of floats between 0 and 1 → use as-is
    if isinstance(color, (list, tuple)) and all(
        isinstance(c, (int, float)) for c in color
    ):
        comps = list(color)

        # If components look like 0-255 integers, scale them down
        if all(isinstance(c, int) and c > 1 for c in comps):
            comps = [c / 255 for c in comps]

        # Ensure exactly 3 components (RGB) – ignore alpha if present, repeat if only one given
        if len(comps) == 1:
            comps = comps * 3
        return tuple(comps[:3])

    # Packed integer 0xRRGGBB
    if isinstance(color, int):
        r = (color >> 16) & 0xFF
        g = (color >> 8) & 0xFF
        b = color & 0xFF
        return (r / 255, g / 255, b / 255)

    # Fallback → black
    return (0.0, 0.0, 0.0)


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------


def get_string_last_char_position_pdf(
    pdf_path: str | Path,
    s: str,
    *,
    position: float | None = None,
) -> float | None:
    """Return the *y* position (distance from the very top of the PDF) of the
    **last character** of the **first** occurrence of *s*.

    When *position* is provided, the function returns the y-coordinate of the
    **first** occurrence whose last character lies *below* (i.e. has a greater
    y-coordinate than) the supplied *position*.  This is useful when multiple
    identical strings exist in the document and you need the match located
    further down the page/document.  If *position* is ``None`` (default), the
    function behaves as before and simply returns the top-most match.

    Parameters
    ----------
    pdf_path: str | pathlib.Path
        Path to the PDF file.
    s: str
        Search string to locate.
    position: float | None, optional
        Reference y-coordinate.  Only matches whose last character baseline is
        *below* this value are considered.  ``None`` means no restriction.

    Returns
    -------
    float | None
        The vertical position in PDF coordinate space (0 at top of first
        page, increasing downwards) of the selected occurrence. ``None`` is
        returned when *s* cannot be found (or no match exists below
        *position* when it is specified).
    """
    if not s:
        raise ValueError("Search string cannot be empty")

    pdf_path = Path(pdf_path)

    with fitz.open(pdf_path) as doc:
        for page_index, page in enumerate(doc):
            page_offset = (
                page_index * page.rect.height
            )  # distance from doc top to this page

            # search_for returns rectangles ordered top-to-bottom
            rects = page.search_for(s, quads=False)
            if not rects:
                continue

            for rect in rects:
                global_y = (
                    page_offset + rect.y1
                )  # y1 = lower bound (last char baseline)

                # Decide whether this match satisfies the *position* criterion
                if position is None or global_y > position:
                    return global_y

    # Not found (either no occurrences at all, or none below *position*)
    return None


def replace_text_preserve_layout_pdf(
    pdf_path: str | Path,
    s1: str | list[str],
    s2: str | list[str],
    out_path: str | Path | None = None,
    *,
    position: float | None = None,
):
    """Replace occurrences of *s1* with *s2* while preserving layout & styling.

    This function now supports replacing **multiple lines** in a single
    operation.  When either *s1* or *s2* contains line-breaks (``"\n"``), the
    strings are split into corresponding *lines* and processed **sequentially
    from top to bottom** within the same opened PDF.  The vertical *position*
    logic that ensures ordering is handled internally, so callers can simply
    provide the full multi-line placeholders and replacements **without**
    needing to pass *position*.

    Backwards-compatibility: For single-line replacements the behaviour is
    identical to previous versions.  The *position* keyword is still accepted
    but ignored when multi-line input is detected.

    Note: The *replacement* text is automatically passed through
    `sanitize_unicode_for_pdf` **after** locating the placeholder so that
    Unicode glyph substitutions cannot break the match while still avoiding
    rendering issues in the written text.
    """
    if not s1:
        raise ValueError("Search string cannot be empty")

    # Determine whether we are in multi-line mode
    multiline = (
        isinstance(s1, list) or isinstance(s2, list) or ("\n" in s1) or ("\n" in s2)
    )

    # Split into lists so that *s1_lines[i]* should be replaced by
    # *s2_lines[i]*.  ``splitlines()`` keeps ordering and drops trailing \n.
    if multiline:
        s1_lines = s1 if isinstance(s1, list) else s1.splitlines()
        s2_lines = s2 if isinstance(s2, list) else s2.splitlines()
    else:
        s1_lines = [s1]
        s2_lines = [s2]

    if len(s1_lines) != len(s2_lines):
        raise ValueError("s1 and s2 must have the same number of lines")

    pdf_path = Path(pdf_path)
    out_path = (
        Path(out_path) if out_path else pdf_path.with_stem(pdf_path.stem + "_filled")
    )

    tmp_dir = Path(tempfile.mkdtemp(prefix="pdf_fonts_"))

    # ------------------------------------------------------------------
    # Local caches to avoid repeated costly operations
    # ------------------------------------------------------------------
    font_cache: dict[str, fitz.Font] = {}

    doc = fitz.open(pdf_path)
    doc_closed = False  # track to avoid double-close

    # Helper used inside the loop -------------------------------------------------
    def _replace_once(
        search: str,
        repl: str,
        last_pos: float | None,
    ) -> float | None:
        """Inner helper that performs **one** find-and-replace and returns
        the vertical position of the *repl*'s last character baseline.  The
        implementation mirrors the previous single-call behaviour but operates
        on the *doc* object that is already open and **does not** save.
        """

        selected_rect: fitz.Rect | None = None
        selected_page_index: int | None = None

        # ------------------------------------------------------------------
        # Locate target rectangle respecting *last_pos* ordering
        # ------------------------------------------------------------------
        for page_index, page in enumerate(doc):
            page_offset = page_index * page.rect.height
            # ``search_for`` may deliver **multiple** rectangles for a single
            # hit (e.g. the phrase spans several words or even lines).  We
            # need all of them to properly redact the old text.  The default
            # behaviour without *hit_max* already returns rectangles grouped
            # by hit and ordered top-to-bottom, so the very first rectangle
            # for a new hit marks the start of that occurrence.
            rects = page.search_for(search, quads=False)
            if not rects:
                continue

            # For cases where multiple rectangles are found (e.g., checkbox + text),
            # we need to find the proper bounding rectangle
            if len(rects) > 1:
                # Look for the line containing the search text
                text_dict = page.get_text("dict")
                for block in text_dict["blocks"]:
                    if block.get("type") != 0:
                        continue
                    for line in block["lines"]:
                        line_text = "".join(span["text"] for span in line["spans"])
                        if search in line_text:
                            line_rect = fitz.Rect(line["bbox"])
                            global_y = page_offset + line_rect.y1

                            if last_pos is None or global_y > last_pos:
                                # Calculate the exact bounds of the search text within the line
                                # by examining the spans
                                search_start_x = None
                                search_end_x = None

                                # Check if we have checkbox-like patterns that might be in separate spans
                                accumulated_text = ""
                                for span in line["spans"]:
                                    span_text = span["text"]
                                    span_bbox = span["bbox"]

                                    if search_start_x is None and search.startswith(
                                        accumulated_text + span_text
                                    ):
                                        search_start_x = span_bbox[0]

                                    accumulated_text += span_text

                                    if accumulated_text.startswith(search):
                                        search_end_x = span_bbox[2]
                                        # We found the complete search text
                                        if len(accumulated_text) >= len(search):
                                            break

                                if (
                                    search_start_x is not None
                                    and search_end_x is not None
                                ):
                                    # Create rectangle that covers the exact search text
                                    selected_rect = fitz.Rect(
                                        search_start_x,
                                        line_rect.y0,
                                        search_end_x,
                                        line_rect.y1,
                                    )
                                    selected_page_index = page_index
                                    break
                    if selected_rect is not None:
                        break

            # ------------------------------------------------------------------
            # Fallback – take the *first* rectangle that lies below *last_pos*
            # and compute the union of *all* rectangles that belong to that
            # same hit so that we blank out the full placeholder.
            # ------------------------------------------------------------------
            if selected_rect is None:
                for hit_start_idx, rect in enumerate(rects):
                    global_y = page_offset + rect.y1  # baseline of this rect

                    if last_pos is not None and global_y <= last_pos:
                        continue

                    # We assume that rectangles belonging to the same hit
                    # are contiguous.  Starting from *hit_start_idx* we
                    # gather all rectangles whose vertical centre is close to
                    # the first one (same line) – this covers the common case
                    # where PyMuPDF splits by spaces.
                    base_y0, base_y1 = rect.y0, rect.y1
                    rects_same_hit: list[fitz.Rect] = [rect]

                    for follower in rects[hit_start_idx + 1 :]:
                        cy = (follower.y0 + follower.y1) / 2
                        if base_y0 - 2 <= cy <= base_y1 + 2:  # small tolerance
                            rects_same_hit.append(follower)
                        else:
                            break  # next hit reached

                    # Union across all rectangles of this hit
                    union_rect = fitz.Rect(rects_same_hit[0])
                    for r_sub in rects_same_hit[1:]:
                        union_rect |= r_sub

                    selected_rect = union_rect
                    selected_page_index = page_index
                    break

            if selected_rect is not None:
                break  # early exit outer loop as soon as match found

        if selected_rect is None:
            logging.debug(
                "No matching occurrences of '%s' found for replacement.", search
            )
            return last_pos  # unchanged

        # Build mapping for this single page only – we store **one**
        # rectangle covering the entire placeholder.  This avoids writing the
        # replacement text multiple times while still ensuring that all
        # original glyphs are removed.
        to_replace: dict[int, list[fitz.Rect]] = {selected_page_index: [selected_rect]}

        # ------------------------------------------------------------------
        # Perform replacement on that page
        # ------------------------------------------------------------------
        for page_index, page in enumerate(doc):
            if page_index not in to_replace:
                continue

            spans: list[dict] = []
            glyphs: list[dict] = []
            for block in page.get_text("rawdict")["blocks"]:
                for line in block.get("lines", []):
                    for sp in line.get("spans", []):
                        spans.append(sp)
                        glyphs.extend(_span_glyphs(sp))

            for rect in to_replace[page_index]:
                r = fitz.Rect(rect)
                target_glyphs = [
                    g for g in glyphs if r.intersects(fitz.Rect(g["bbox"]))
                ]
                if not target_glyphs:
                    continue

                # Sanitize replacement text only **after** locating the placeholder so
                # that Unicode substitutions cannot break the search phase.
                sanitized_repl = sanitize_unicode_for_pdf(repl)

                # ----------------------------------------------------------
                # Dynamically determine a baseline that best matches the
                # original line.  *search_for* often returns glyph boxes that
                # include a small descent area; using bbox[3] (bottom) as the
                # baseline therefore puts freshly written text a tad too low.
                #
                # Approach:
                # 1. Collect vertical metrics from all glyphs that belong to
                #    the placeholder – this gives us a representative min top
                #    and max bottom.
                # 2. Derive the glyph height and compare it to the reported
                #    *font size* – any excess is assumed to be descent /
                #    leading that we want to subtract partly.
                # 3. Move the baseline *up* by twice the excess (capped at
                #    20 % of the font size) so the new text sits visually on
                #    the same line as the deleted placeholder.
                # 4. Special handling for checkbox characters which have different
                #    vertical metrics than regular text.

                # Define checkbox characters
                checkbox_chars = {"□", "■", "☐", "☑", "☒", "⬜", "⬛"}

                # Separate checkbox glyphs from text glyphs
                text_glyphs = []
                checkbox_glyphs = []
                for g in target_glyphs:
                    if g.get("c") in checkbox_chars:
                        checkbox_glyphs.append(g)
                    else:
                        text_glyphs.append(g)

                # Use text glyphs for baseline calculation if available,
                # otherwise fall back to all glyphs
                baseline_glyphs = text_glyphs if text_glyphs else target_glyphs

                base_glyph = baseline_glyphs[0]
                fontsize = base_glyph["size"]

                glyph_bottoms = [g["bbox"][3] for g in baseline_glyphs]
                glyph_tops = [g["bbox"][1] for g in baseline_glyphs]

                baseline_raw = max(glyph_bottoms)
                glyph_top = min(glyph_tops)
                glyph_height = baseline_raw - glyph_top

                excess = max(0, glyph_height - fontsize)  # space below baseline

                # For text with checkboxes, use less aggressive baseline shift
                if checkbox_glyphs and text_glyphs:
                    # Use minimal shift for checkbox+text combinations
                    baseline_shift = min(excess * 0.5, fontsize * 0.1)
                else:
                    # Original calculation for regular text
                    baseline_shift = min(excess * 1.8, fontsize * 0.2)

                baseline_y = baseline_raw - baseline_shift
                baseline = fitz.Point(rect.x0, baseline_y)

                # redact placeholder text
                page.add_redact_annot(rect, fill=(1, 1, 1))
                page.apply_redactions()

                cursor_x = baseline.x
                writer = fitz.TextWriter(page.rect)

                def _get_font(name: str):
                    if name in font_cache:
                        return font_cache[name]

                    # 1. Try embedded font by name
                    try:
                        font_cache[name] = fitz.Font(fontname=name)
                        return font_cache[name]
                    except Exception:
                        pass

                    # 2. Try PyMuPDF's built-in font mapping
                    # Map common PDF fonts to PyMuPDF's base 14 fonts
                    pymupdf_font_map = {
                        "lmroman": "tiro",  # Times-like serif font
                        "cmr": "tiro",  # Computer Modern Roman
                        "times": "tiro",
                        # Map sans-serif fonts to Helvetica regular variant
                        "arial": "helv",  # Helvetica-like sans-serif (regular)
                        "helvetica": "helv",
                        "lmsans": "helv",  # Latin Modern Sans
                        "cmss": "helv",  # Computer Modern Sans
                        "courier": "cour",  # Courier
                        "lmmono": "cour",  # Latin Modern Mono
                        "cmtt": "cour",  # Computer Modern Typewriter
                    }

                    # Try to match font name prefix
                    name_lower = name.lower()
                    for prefix, pymupdf_font in pymupdf_font_map.items():
                        if name_lower.startswith(prefix):
                            try:
                                font_cache[name] = fitz.Font(fontname=pymupdf_font)
                                logging.info(
                                    f"Font mapping: {name} → {pymupdf_font} (PyMuPDF built-in)"
                                )
                                return font_cache[name]
                            except Exception:
                                pass
                    # 3. Resolve via font_manager helper
                    resolved_name, font_path = get_available_font(
                        name, cache_dir=str(tmp_dir)
                    )

                    if font_path:
                        try:
                            font_cache[name] = fitz.Font(fontfile=str(font_path))
                            return font_cache[name]
                        except Exception:
                            logging.debug(
                                "Failed to embed font file %s, falling back to font name.",
                                font_path,
                            )

                    if (
                        resolved_name.lower() == "times new roman"
                        or "times" in resolved_name.lower()
                    ):
                        try:
                            font_cache[name] = fitz.Font(fontname="tiro")
                            return font_cache[name]
                        except:
                            pass
                    elif (
                        resolved_name.lower() == "arial"
                        or "arial" in resolved_name.lower()
                        or resolved_name.lower() == "helvetica"
                        or "helvetica" in resolved_name.lower()
                    ):
                        try:
                            font_cache[name] = fitz.Font(fontname="helv")
                            return font_cache[name]
                        except:
                            pass

                    try:
                        font_cache[name] = fitz.Font(fontname=resolved_name)
                    except Exception:
                        font_cache[name] = fitz.Font(fontname="helv")
                    return font_cache[name]

                # ------------------------------------------------------------------
                # Determine the dominant font (most frequently occurring) among the
                # glyphs that constitute the placeholder so that the replacement text
                # visually matches the majority of the original content.  Fall back to
                # the first glyph's font if the Counter is empty for any reason.
                # ------------------------------------------------------------------

                font_counter = Counter(
                    g["font"] for g in target_glyphs if g.get("font")
                )
                dominant_fontname = (
                    font_counter.most_common(1)[0][0]
                    if font_counter
                    else target_glyphs[0]["font"]
                )

                # Find a representative glyph that uses that dominant font to copy
                # the size and colour from. If none is found (edge-case) we again fall
                # back to the first glyph.
                rep_glyph = next(
                    (g for g in target_glyphs if g.get("font") == dominant_fontname),
                    target_glyphs[0],
                )

                fontsize = rep_glyph["size"]
                rep_color = _normalize_color(rep_glyph["color"])

                font_obj = _get_font(dominant_fontname)

                # Apply colour once at the start – replacement string is assumed to
                # use a single colour.
                writer.color = rep_color

                for ch in sanitized_repl:
                    glyph_w = font_obj.text_length(ch, fontsize)

                    writer.append(
                        fitz.Point(cursor_x, baseline.y),
                        ch,
                        font_obj,
                        fontsize,
                    )
                    cursor_x += glyph_w

                writer.write_text(page, overlay=True)

        # Return the baseline y of the replacement for ordering of next one
        page_offset = selected_page_index * doc[selected_page_index].rect.height
        return page_offset + selected_rect.y1

    # ------------------------------------------------------------------
    # Sequential processing – respects vertical ordering automatically.
    # We now *honour* the caller-supplied *position* even in multi-line mode so
    # that the very first search begins below that reference.  Subsequent
    # replacements update *last_position* incrementally.
    # ------------------------------------------------------------------
    n = len(s1_lines)
    related_lines = {s1_lines[i] for i in range(n) if s1_lines[i] != s2_lines[i]}
    cleaned_s1_lines = []
    cleaned_s2_lines = []
    for i in range(n):
        if s1_lines[i] in related_lines:
            cleaned_s1_lines.append(s1_lines[i])
            cleaned_s2_lines.append(s2_lines[i])
    s1_lines = cleaned_s1_lines
    s2_lines = cleaned_s2_lines

    last_position: float | None = position  # always start from the caller hint

    for search_line, repl_line in zip(s1_lines, s2_lines):
        # Skip completely empty placeholders – these are artefacts of splitlines()
        if not search_line:
            continue

        if search_line == repl_line:
            last_position = get_string_last_char_position_pdf(
                pdf_path, search_line, position=last_position
            )
        else:
            last_position = _replace_once(search_line, repl_line, last_position)

    # ------------------------------------------------------------------
    # Save / persist changes exactly once
    # ------------------------------------------------------------------
    try:
        if out_path.resolve() == pdf_path.resolve():
            tmp_out = out_path.with_suffix(".tmp.pdf")
            doc.save(tmp_out, deflate=True)
            doc.close()
            doc_closed = True
            shutil.move(tmp_out, out_path)
        else:
            doc.save(out_path, deflate=True)
            doc.close()
            doc_closed = True
    finally:
        if not doc_closed:
            try:
                doc.close()
            except Exception:
                pass
        shutil.rmtree(tmp_dir, ignore_errors=True)

    return out_path


def _iter_paragraphs_with_index(doc: Document):
    """Yield paragraphs from a python-docx Document in visual order.

    Returns tuples ``(paragraph, global_index)`` where *global_index* is the
    zero-based position counting all paragraphs including those inside tables
    exactly at the spot where the table appears.
    """
    index = 0
    for block in doc.element.body.iterchildren():
        if block.tag.endswith("}p"):
            yield (Paragraph(block, doc), index)
            index += 1
        elif block.tag.endswith("}tbl"):
            tbl = block
            for row in tbl.iterchildren():
                if not row.tag.endswith("}tr"):
                    continue
                for cell in row.iterchildren():
                    if not cell.tag.endswith("}tc"):
                        continue
                    for cell_block in cell.iterchildren():
                        if cell_block.tag.endswith("}p"):
                            yield (Paragraph(cell_block, doc), index)
                            index += 1
    if index == 0:
        # Fallback for unusual docs – just rely on Document.paragraphs order
        for para in doc.paragraphs:
            yield (para, index)
            index += 1


def get_string_last_char_position_docx(
    docx_path: str | Path,
    s: str,
    *,
    position: tuple[int, int] | None = None,
) -> tuple[int, int] | None:
    """Return a *location* tuple ``(paragraph_index, char_index)`` of the **last
    character** of the **first** occurrence of *s* in a DOCX file.

    The returned position can later be fed back via the *position* parameter to
    continue searches *below* that reference – similar to the PDF helper that
    operates on y-coordinates.  Ordering is determined by paragraph order in
    the document (including paragraphs in tables) and the character offset
    inside that paragraph.
    """
    if not s:
        raise ValueError("Search string cannot be empty")

    docx_path = Path(docx_path)
    doc = Document(docx_path)

    # Perform search
    for para, idx in _iter_paragraphs_with_index(doc):
        text = para.text
        if not text:
            continue
        found_at = text.find(s)
        if found_at == -1:
            continue
        last_char_idx = found_at + len(s) - 1

        if (
            position is None
            or (idx > position[0])
            or (idx == position[0] and last_char_idx > position[1])
        ):
            return (idx, last_char_idx)

    return None


def replace_text_preserve_layout_docx(
    docx_path: str | Path,
    s1: str | list[str],
    s2: str | list[str],
    out_path: str | Path | None = None,
    *,
    position: tuple[int, int] | None = None,
):
    """Replace *s1* with *s2* in a DOCX file while attempting to preserve layout.

    The function mirrors the public interface of the PDF variant – including
    support for multi-line placeholders and the *position* parameter that
    allows sequential replacements from top-to-bottom.  Styling is preserved
    to the extent that we modify text **within existing runs** instead of
    recreating paragraphs from scratch.  This is generally sufficient for the
    placeholders generated by the form-filling pipeline which reside entirely
    within a single run.
    """
    if not s1:
        raise ValueError("Search string cannot be empty")

    # Determine whether we are in multi-line mode
    multiline = (
        isinstance(s1, list) or isinstance(s2, list) or ("\n" in s1) or ("\n" in s2)
    )

    # Split into lists so that *s1_lines[i]* should be replaced by
    # *s2_lines[i]*.  ``splitlines()`` keeps ordering and drops trailing \n.
    if multiline:
        s1_lines = s1 if isinstance(s1, list) else s1.splitlines()
        s2_lines = s2 if isinstance(s2, list) else s2.splitlines()
    else:
        s1_lines = [s1]
        s2_lines = [s2]

    if len(s1_lines) != len(s2_lines):
        raise ValueError("s1 and s2 must have the same number of lines")

    docx_path = Path(docx_path)
    out_path = (
        Path(out_path) if out_path else docx_path.with_stem(docx_path.stem + "_filled")
    )

    # Open once and keep open for sequential replacements
    doc = Document(docx_path)

    # ----------------- helper for single replacement --------------------
    def _replace_once(search: str, repl: str, last_pos: tuple[int, int] | None):
        """Replace **one** occurrence of *search* that lies *after* *last_pos*.

        Returns the position tuple of the last char of the *repl* after
        replacement so that subsequent calls can continue further down.
        """
        # Iterate through paragraphs in order, keeping index
        # We'll reuse the generator defined earlier
        for para, idx in _iter_paragraphs_with_index(doc):
            text = para.text
            if not text:
                continue
            # ----------------------- modified search logic -----------------------
            # Find the first occurrence of the placeholder *after* last_pos.
            if last_pos is not None:
                # Skip paragraphs that lie completely above last_pos
                if idx < last_pos[0]:
                    continue

                # Determine the offset inside the paragraph where searching should begin
                search_from = 0 if idx > last_pos[0] else last_pos[1]
                start = text.find(search, search_from)
            else:
                start = text.find(search)

            # No occurrence in this paragraph that satisfies the ordering constraint
            if start == -1:
                continue

            end_char = start + len(search)

            para.text = para.text[:start] + repl + para.text[end_char:]
            return (idx, start + len(repl))

        # No match
        logging.debug("No matching occurrences of '%s' found for replacement.", search)
        return last_pos

    n = len(s1_lines)
    related_lines = {s1_lines[i] for i in range(n) if s1_lines[i] != s2_lines[i]}
    cleaned_s1_lines = []
    cleaned_s2_lines = []
    for i in range(n):
        if s1_lines[i] in related_lines:
            cleaned_s1_lines.append(s1_lines[i])
            cleaned_s2_lines.append(s2_lines[i])
    s1_lines = cleaned_s1_lines
    s2_lines = cleaned_s2_lines

    # Sequential processing ------------------------------------------------------
    last_position: tuple[int, int] | None = position
    for search_line, repl_line in zip(s1_lines, s2_lines):
        if not search_line:
            continue
        if search_line == repl_line:
            last_position = get_string_last_char_position_docx(
                docx_path, search_line, position=last_position
            )
            continue
        last_position = _replace_once(search_line, repl_line, last_position)

    # Save -----------------------------------------------------------------------
    doc.save(out_path)
    return out_path


def extract_text_with_fields_as_underscores(
    pdf_path: str | Path,
) -> tuple[str, list[tuple[int, str]]]:
    """
    Extract text from a PDF in reading order, replacing text field positions with underscores.

    This function processes a PDF and returns its text content where any interactive
    text fields (form fields) are replaced with "_______" at their original positions,
    along with field metadata.

    Parameters
    ----------
    pdf_path : str | Path
        Path to the PDF file to process.

    Returns
    -------
    tuple[str, list[tuple[int, str]]]
        A tuple containing:
        - The extracted text with text fields replaced by underscores
        - A list of tuples (field_flags, field_name) for each text field in reading order

    Example
    -------
    If a PDF contains "First name [text_field]", the output will be:
    ("First name _______", [(0, "first_name_field")])
    """
    pdf_path = Path(pdf_path)

    with fitz.open(pdf_path) as doc:
        full_text = []
        field_info_list = []  # To store (field_flags, field_name) in reading order

        for page_idx, page in enumerate(doc):
            # Get all text with position information
            text_dict = page.get_text("dict")

            # Get all form fields (widgets) on this page
            text_fields = []
            for widget in page.widgets():
                if widget.field_type == fitz.PDF_WIDGET_TYPE_TEXT:  # Text fields only
                    text_fields.append(
                        {
                            "type": "field",
                            "x": widget.rect.x0,
                            "y": widget.rect.y0,
                            "text": "_______",
                            "rect": widget.rect,
                            "field_flags": widget.field_flags,
                            "field_name": widget.field_name,
                        }
                    )

            # Collect all text elements with their positions
            text_elements = []

            for block in text_dict["blocks"]:
                if block.get("type") != 0:  # Skip non-text blocks
                    continue

                for line in block["lines"]:
                    for span in line["spans"]:
                        if span["text"].strip():  # Skip empty text
                            text_elements.append(
                                {
                                    "type": "text",
                                    "x": span["bbox"][0],
                                    "y": span["bbox"][1],
                                    "text": span["text"],
                                    "bbox": span["bbox"],
                                }
                            )

            # Combine text elements and fields
            all_elements = text_elements + text_fields

            # Sort by Y coordinate (top to bottom), then by X coordinate (left to right)
            all_elements.sort(key=lambda elem: (elem["y"], elem["x"]))

            # Group elements into lines based on Y coordinate
            lines = []
            if all_elements:
                # First pass: group elements by similar Y coordinates
                line_groups = []
                current_group = [all_elements[0]]

                for elem in all_elements[1:]:
                    # Check if this element belongs to the current line group
                    # Use the average Y of the current group for comparison
                    avg_y = sum(e["y"] for e in current_group) / len(current_group)

                    # More sophisticated line detection: consider both Y distance and element types
                    y_threshold = 5 if elem["type"] == "field" else 3

                    if abs(elem["y"] - avg_y) < y_threshold:
                        current_group.append(elem)
                    else:
                        # Save current group and start a new one
                        line_groups.append(current_group)
                        current_group = [elem]

                # Don't forget the last group
                if current_group:
                    line_groups.append(current_group)

                # Second pass: process each line group
                for group in line_groups:
                    # Sort elements in the group by X coordinate
                    group.sort(key=lambda e: e["x"])

                    # Build the line text with proper spacing
                    line_parts = []
                    prev_end_x = None

                    for element in group:
                        # Calculate spacing based on the gap between elements
                        if prev_end_x is not None:
                            gap = element["x"] - prev_end_x

                            # Add appropriate spacing
                            if gap > 50:  # Large gap
                                line_parts.append("    ")  # Multiple spaces
                            elif gap > 20:  # Medium gap
                                line_parts.append("  ")  # Double space
                            elif gap > 5:  # Small gap
                                line_parts.append(" ")  # Single space
                            # If gap is very small or negative, don't add space (text might be touching)

                        # Add the element text
                        line_parts.append(element["text"].rstrip())

                        # If this is a field, record its info in reading order
                        if element["type"] == "field":
                            field_info_list.append(
                                (element["field_flags"], element["field_name"])
                            )

                        # Update prev_end_x for next iteration
                        if element["type"] == "text" and "bbox" in element:
                            # For text, use actual bounding box to calculate end position
                            prev_end_x = element["bbox"][2]  # x1 coordinate
                        else:
                            # For fields, estimate based on the underscore length
                            prev_end_x = element["x"] + 50  # Approximate field width

                    # Join all parts and clean up the line
                    line_text = "".join(line_parts).strip()

                    # Only add non-empty lines
                    if line_text:
                        # Check if this line contains only field markers and should be merged with previous line
                        if (
                            lines
                            and line_text.strip() == "_______"
                            and group[0]["type"] == "field"
                        ):
                            # This is likely a field that belongs to the previous line
                            # Check if previous line ends with text that suggests a field should follow
                            prev_line = lines[-1]
                            if (
                                prev_line.endswith(")")
                                or prev_line.endswith(":")
                                or prev_line.endswith("by")
                                or prev_line.endswith("address")
                                or prev_line.endswith("Address")
                                or not prev_line.endswith("_______")
                            ):
                                # Merge with previous line
                                lines[-1] = prev_line + " " + line_text
                            else:
                                lines.append(line_text)
                        else:
                            lines.append(line_text)

            # Add page text to full text
            if lines:
                page_text = "\n".join(lines)
                full_text.append(page_text)

        # Join all pages with double newlines
        return "\n\n".join(full_text), field_info_list


def extract_form_text(form_path: str) -> str:
    """Extract all text from a form file (DOCX or PDF) for pattern analysis."""
    ext = os.path.splitext(form_path)[1].lower()

    if ext == ".docx":
        doc = Document(form_path)
        lines = []

        # Extract from paragraphs
        for para in doc.paragraphs:
            lines.append(para.text)

        # Extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        lines.append(para.text)

        return "\n".join(lines)
    elif ext == ".pdf":
        return extract_text_with_fields_as_underscores(form_path)[0]
    else:
        raise ValueError(f"Unsupported form format: {ext}")


# ---------------------------------------------------------------------------
# DOCX → PDF conversion helper
# ---------------------------------------------------------------------------

def convert_docx_to_pdf(docx_path: str | Path, *, remove_source: bool = False) -> Path:
    """Convert a DOCX file to PDF using *LibreOffice* in headless mode.

    This utility calls ``soffice --headless --convert-to pdf`` which is
    available once LibreOffice is installed and added to the system *PATH*.
    On success the generated PDF sits in the same directory as the source file
    and the absolute :class:`pathlib.Path` to that PDF is returned.

    Parameters
    ----------
    docx_path:
        Path to the input *.docx* file.
    remove_source:
        When *True* the original DOCX file is deleted after a successful
        conversion.

    Returns
    -------
    pathlib.Path
        Path to the generated PDF file.
    """

    from pathlib import Path  # local import
    import shutil
    import subprocess

    docx_path = Path(docx_path).expanduser().resolve()
    if not docx_path.exists():
        raise FileNotFoundError(docx_path)

    soffice = shutil.which("soffice")
    if soffice is None:
        raise RuntimeError(
            "LibreOffice (soffice) not found on PATH. Install LibreOffice to enable DOCX → PDF conversion."
        )

    cmd = [
        soffice,
        "--headless",
        "--convert-to",
        "pdf",
        "--outdir",
        str(docx_path.parent),
        str(docx_path),
    ]
    subprocess.run(cmd, check=True)

    pdf_path = docx_path.with_suffix(".pdf")
    if not pdf_path.exists():
        raise RuntimeError("LibreOffice did not produce the expected PDF file")

    if remove_source:
        try:
            docx_path.unlink(missing_ok=True)
        except Exception:
            logging.debug("Could not delete temporary DOCX %s", docx_path)

    return pdf_path


# ---------------------------------------------------------------------------
# PDF → DOCX conversion helper
# ---------------------------------------------------------------------------

def convert_pdf_to_docx(
    pdf_path: str | Path,
    docx_path: str | Path | None = None,
    *,
    start: int = 0,
    end: int | None = None,
    multi_processing: bool = True,
) -> Path:
    """Convert a PDF file to a DOCX document.

    This is a thin wrapper around *pdf2docx*'s :class:`pdf2docx.Converter` that
    exposes the most common options with sensible defaults and integrates with
    the rest of *text_utils* by adopting ``pathlib.Path`` objects and returning
    the destination path.

    Parameters
    ----------
    pdf_path:
        Path to the source PDF file.
    docx_path:
        Output path for the generated *.docx* file.  When *None* (default) the
        function writes alongside *pdf_path* using the same stem.
    start:
        First page (zero-based) to convert – defaults to the beginning.
    end:
        Last page (inclusive) to convert – *None* converts until the end.
    multi_processing:
        Whether to enable *pdf2docx*'s multi-processing mode which can speed up
        large conversions at the cost of additional memory.

    Returns
    -------
    pathlib.Path
        The absolute path to the created DOCX file.
    """

    from pathlib import Path  # local import to keep global namespace clean

    try:
        from pdf2docx import Converter  # type: ignore
    except ImportError as exc:
        raise ImportError(
            "pdf2docx is required for PDF-to-DOCX conversion. Install it via 'pip install pdf2docx'."
        ) from exc

    pdf_path = Path(pdf_path)
    if docx_path is None:
        docx_path = pdf_path.with_suffix(".docx")
    else:
        docx_path = Path(docx_path)

    # Perform the conversion – wrap in try/finally to ensure the converter is closed
    converter = Converter(str(pdf_path))
    try:
        converter.convert(
            str(docx_path),
            start=start,
            end=end,
            multi_processing=multi_processing,
        )
    finally:
        converter.close()

    logging.debug("Converted '%s' → '%s' (pages %s-%s)", pdf_path, docx_path, start, end or "end")
    return docx_path



def file_checksum(path: str) -> str:
    """Return a short SHA-256 checksum based on *path* and its last-modified time.

    Combining the absolute path with the file's mtime means the checksum
    changes whenever the file is edited, preserving the previous behaviour
    where the cache invalidates on modification time changes.
    """
    mtime = os.path.getmtime(path)
    data = f"{os.path.abspath(path)}:{mtime}".encode("utf-8")
    return hashlib.sha256(data).hexdigest()[:16]


if __name__ == "__main__":
    # Example usage
    text, field_info = extract_text_with_fields_as_underscores(
        "/Users/hieu.leduc/Desktop/easyform/forms/application1.pdf"
    )

    print("=== EXTRACTED TEXT ===")
    print(text[:1000] + "\n... [truncated] ...")

    print("\n=== FIELD INFORMATION (in reading order) ===")
    print(f"Total fields: {len(field_info)}")
    for i, (flags, name) in enumerate(field_info[:10]):  # Show first 10 fields
        print(f"{i+1}. Field name: {name}, Flags: {flags}")
    if len(field_info) > 10:
        print(f"... and {len(field_info) - 10} more fields")