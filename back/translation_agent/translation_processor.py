import os
import logging
from typing import List
from pydantic import ValidationError

# Additional dependency for DOCX manipulation
from docx import Document

from ..filler_agent.text_utils import file_checksum, convert_pdf_to_docx, convert_docx_to_pdf
from ..llm_client import query_gpt
from ..llm_priority_queue import Priority

# Extract helpers
from ..context_extraction.context_extractor import (
    extract_pdf,
    extract_text,
    extract_image,
)
from .prompts import TRANSLATION_PROMPT_TEMPLATE
from .prompts import PARAGRAPH_TRANSLATION_PROMPT_TEMPLATE
from .output_schemas import MarkdownTranslationModel, ParagraphTranslationModel


# -----------------------------------------------------------------------------
# Core translation logic
# -----------------------------------------------------------------------------

CACHED_TRANSLATIONS = {}


def translate_markdown(
    markdown: str,
    target_lang: str,
    provider: str = "openai",
    *,
    cancel_id: str | None = None,
    from_chat: bool = False,
) -> str:
    """Translate a Markdown document string into *target_lang*.

    Parameters
    ----------
    markdown: str
        The original Markdown content.
    target_lang: str
        The language to translate the content into (e.g. "fr", "de", "Vietnamese").
    provider: str, optional
        The LLM provider identifier understood by ``llm_client.query_gpt``.

    Returns
    -------
    str
        The translated Markdown string produced by the LLM.
    """
    prompt = TRANSLATION_PROMPT_TEMPLATE.format(language=target_lang, markdown=markdown)
    raw_resp = query_gpt(
        prompt,
        provider=provider,
        cancel_id=cancel_id,
        response_format=MarkdownTranslationModel,
        priority=Priority.NORMAL if not from_chat else Priority.URGENT,
    )

    # The client returns raw JSON string adhering to MarkdownTranslationModel
    try:
        obj = MarkdownTranslationModel.model_validate_json(raw_resp)
    except (ValidationError, ValueError):
        # final fallback – try to parse manually
        import re

        try:
            cleaned = raw_resp.strip().strip("` ")
            # grab first json object
            m = re.search(r"\{.*\}", cleaned, re.DOTALL)
            if m:
                obj = MarkdownTranslationModel.model_validate_json(m.group(0))
            else:
                raise
        except Exception as exc:
            logging.error(
                f"Failed to validate translation response: {exc}\nRaw response: {raw_resp[:200]}…"
            )
            raise RuntimeError("LLM did not return valid translation JSON") from exc
    return obj.markdown


# -----------------------------------------------------------------------------
# New: Paragraph container for DOCX-specific translation
# -----------------------------------------------------------------------------


class Paragraph:
    """Represents a DOCX paragraph and its translation."""

    def __init__(self, original_text: str):
        self.original_text: str = original_text
        self.translated_text: str = ""

    def __repr__(self) -> str:  # pragma: no cover
        return f"Paragraph(orig={self.original_text!r}, trans={self.translated_text!r})"


def has_text_layer(pdf_path: str) -> bool:
    """Return True if the PDF at *pdf_path* contains any extractable text.

    The function first tries to use the PyMuPDF (``fitz``) backend. If that is
    not available it falls back to *pypdf*. We iterate through the pages and
    return as soon as we encounter any non-whitespace characters.  If neither
    backend is present or text extraction fails the function conservatively
    returns *False*.
    """

    # ------------------------------------------------------------------
    # Try PyMuPDF first – fastest and most reliable for text layer detection
    # ------------------------------------------------------------------
    try:
        import fitz  # type: ignore
    except ImportError:
        fitz = None  # type: ignore

    if fitz is not None:
        try:
            with fitz.open(pdf_path) as doc:
                for page in doc:  # type: ignore[assignment]
                    if page.get_text("text").strip():
                        return True
        except Exception as exc:  # pragma: no cover – any extraction failure
            logging.debug(f"PyMuPDF text detection failed for {pdf_path}: {exc}")

    # ------------------------------------------------------------------
    # Fallback: use pypdf (pure-Python) if available
    # ------------------------------------------------------------------
    try:
        from pypdf import PdfReader  # type: ignore
    except ImportError:
        PdfReader = None  # type: ignore

    if PdfReader is not None:
        try:
            reader = PdfReader(pdf_path)
            for page in reader.pages:  # type: ignore[attr-defined]
                text = page.extract_text()
                if text and text.strip():
                    return True
        except Exception as exc:  # pragma: no cover
            logging.debug(f"pypdf text detection failed for {pdf_path}: {exc}")

    # If we reach here we did not find any text
    return False


def translate_file(
    path: str,
    target_lang: str,
    provider: str = "openai",
    output_path: str | None = None,
) -> str:
    """Translate *path* (PDF or DOCX) to *target_lang* and persist next to original.

    Returns the path to the translated document.
    """
    if not os.path.isfile(path):
        raise FileNotFoundError(path)

    # ------------------------------------------------------------------
    # Cache key – use checksum instead of (path, mtime)
    # ------------------------------------------------------------------
    key = (file_checksum(path), target_lang, provider)
    if key in CACHED_TRANSLATIONS:
        logging.info(f"Using cached translation for {path}")
        output_path = CACHED_TRANSLATIONS[key]
        if os.path.isfile(output_path):
            return output_path
        else:
            logging.info(f"Cached translation for {path} is not a file, re-translating")

    _, ext = os.path.splitext(path)
    if ext.lower() not in {".pdf", ".docx", ".png", ".jpg", ".jpeg", ".tiff"}:
        raise ValueError("Only PDF, DOCX, or image inputs are supported")

    # ------------------------------------------------------------------
    # Determine default output path using checksum-based filename
    # ------------------------------------------------------------------
    if output_path is None:
        # Create persistent translation directory in user's home
        translation_dir = os.path.expanduser("~/.EasyFormTranslation")
        os.makedirs(translation_dir, exist_ok=True)

        # Short checksum of the absolute file path – keeps name deterministic but short
        checksum = file_checksum(path)

        # Decide extension for translated output
        if ext.lower() == ".docx":
            out_ext = ".docx"
        else:
            out_ext = ".pdf" if has_text_layer(path) else ".md"

        output_path = os.path.join(
            translation_dir, f"{checksum}_{target_lang}{out_ext}"
        )
        if os.path.isfile(output_path):
            logging.info(
                f"Translation output already exists at {output_path}, skipping translation."
            )
            CACHED_TRANSLATIONS[key] = output_path
            return output_path

    # ------------------------------------------------------------------
    # PDF → DOCX → PDF workflow
    # ------------------------------------------------------------------
    if ext.lower() == ".pdf":
        if has_text_layer(path):
            if not output_path.lower().endswith(".pdf"):
                output_path = os.path.splitext(output_path)[0] + ".pdf"
            final_pdf_path = output_path
            try:
                # 1) Convert PDF to a temporary DOCX
                tmp_docx_path = convert_pdf_to_docx(path)

                # 2) Translate that DOCX (produces another DOCX)
                translated_docx_path = _translate_docx(tmp_docx_path, target_lang, provider)

                # 3) Convert translated DOCX → final PDF
                final_pdf_path = convert_docx_to_pdf(translated_docx_path)

                if output_path is not None:
                    if not output_path.lower().endswith(".pdf"):
                        output_path = os.path.splitext(output_path)[0] + ".pdf"
                    import shutil
                    shutil.move(final_pdf_path, output_path)
                    final_pdf_path = output_path

                # Clean up temporary DOCX files
                try:
                    os.remove(tmp_docx_path)
                except Exception:
                    logging.debug("Could not delete temporary file %s", tmp_docx_path)
                try:
                    os.remove(translated_docx_path)
                except Exception:
                    logging.debug("Could not delete translated DOCX %s", translated_docx_path)
            except Exception as e:
                markdown_text, _ = extract_pdf(path, True)
                translated_md = translate_markdown(markdown_text, target_lang, provider, cancel_id=path)
                final_pdf_path = output_path.replace(".pdf", ".md")
                with open(final_pdf_path, "w", encoding="utf-8") as fh:
                    fh.write(translated_md)
                logging.info(f"Translated markdown written to {final_pdf_path}")
            CACHED_TRANSLATIONS[key] = str(final_pdf_path)
            return str(final_pdf_path)
        else:
            markdown_text, _ = extract_pdf(path, True)
            translated_md = translate_markdown(markdown_text, target_lang, provider, cancel_id=path)
            with open(output_path, "w", encoding="utf-8") as fh:
                fh.write(translated_md)
            logging.info(f"Translated markdown written to {output_path}")
            CACHED_TRANSLATIONS[key] = output_path
            return output_path
    else:
        # DOCX advanced
        output_path = _translate_docx(path, target_lang, provider, output_path)
        CACHED_TRANSLATIONS[key] = output_path
        return output_path


# -----------------------------------------------------------------------------
# Advanced DOCX translation (per-paragraph with context window)
# -----------------------------------------------------------------------------


def _translate_paragraph(
    current: str,
    prev_translated: str,
    next_original: str,
    lang: str,
    provider: str,
    *,
    cancel_id: str | None = None,
) -> str:
    """Query LLM to translate a single paragraph with surrounding context."""
    # If context is missing, explicitly mention it in the prompt so the LLM
    # is aware of document boundaries.
    if not prev_translated.strip():
        prev_section = (
            "[NONE] – *current paragraph is the FIRST paragraph of the document*."
        )
    else:
        prev_section = prev_translated

    if not next_original.strip():
        next_section = (
            "[NONE] – *current paragraph is the LAST paragraph of the document*."
        )
    else:
        next_section = next_original

    prompt = PARAGRAPH_TRANSLATION_PROMPT_TEMPLATE.format(
        language=lang,
        prev_translated=prev_section,
        next_original=next_section,
        current=current,
    )

    raw_resp = query_gpt(
        prompt,
        provider=provider,
        cancel_id=cancel_id,
        response_format=ParagraphTranslationModel,
        priority=Priority.HIGH,
    )

    try:
        obj = ParagraphTranslationModel.model_validate_json(raw_resp)
    except (ValidationError, ValueError):
        # relaxed fallback similar to earlier
        import re

        m = re.search(r"\{.*\}", raw_resp, re.DOTALL)
        if not m:
            logging.error("Paragraph translation response missing JSON object")
            raise RuntimeError("Invalid paragraph translation response")
        obj = ParagraphTranslationModel.model_validate_json(m.group(0))

    return obj.text


def _translate_docx(
    path: str, target_lang: str, provider: str, output_path: str | None = None
) -> str:
    """Translate DOCX file paragraph-by-paragraph and save translated DOCX."""

    doc = Document(path)
    paragraph_objs: List[Paragraph] = [Paragraph(p.text) for p in doc.paragraphs]

    n = len(paragraph_objs)
    for i in range(n):
        cur_para = paragraph_objs[i]
        if not cur_para.original_text.strip():
            # blank paragraph – copy as-is
            cur_para.translated_text = ""
            continue

        prev_translated = "" if i == 0 else paragraph_objs[i - 1].translated_text
        next_original = "" if i == n - 1 else paragraph_objs[i + 1].original_text

        logging.debug(f"Translating paragraph {i+1}/{n}")
        cur_para.translated_text = _translate_paragraph(
            cur_para.original_text,
            prev_translated,
            next_original,
            target_lang,
            provider,
            cancel_id=path,
        )

    # Write back translations to the Document object
    for i, p in enumerate(doc.paragraphs):
        p.text = paragraph_objs[i].translated_text

    # determine output path
    if output_path is None:
        base, _ = os.path.splitext(path)
        output_path = f"{base}_translated_{target_lang}.docx"

    doc.save(output_path)
    logging.info(f"Translated DOCX written to {output_path}")
    return output_path
